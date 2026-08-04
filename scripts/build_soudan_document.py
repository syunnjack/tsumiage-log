from pathlib import Path

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "output" / "docs" / "saimuseiri_soudan_memo_google_docs.docx"


def set_run_font(run, size=11, bold=False, color="000000"):
    run.font.name = "Arial"
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.color.rgb = RGBColor.from_string(color)
    rpr = run._element.get_or_add_rPr()
    rfonts = rpr.rFonts
    if rfonts is None:
        rfonts = OxmlElement("w:rFonts")
        rpr.insert(0, rfonts)
    for key in ("ascii", "hAnsi", "eastAsia"):
        rfonts.set(qn(f"w:{key}"), "Arial")


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for m, v in (("top", top), ("start", start), ("bottom", bottom), ("end", end)):
        node = tc_mar.find(qn(f"w:{m}"))
        if node is None:
            node = OxmlElement(f"w:{m}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(v))
        node.set(qn("w:type"), "dxa")


def set_table_geometry(table, widths):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    tbl_pr = table._tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    tbl_w.set(qn("w:w"), str(sum(widths)))
    tbl_w.set(qn("w:type"), "dxa")
    tbl_ind = tbl_pr.first_child_found_in("w:tblInd")
    if tbl_ind is None:
        tbl_ind = OxmlElement("w:tblInd")
        tbl_pr.append(tbl_ind)
    tbl_ind.set(qn("w:w"), "0")
    tbl_ind.set(qn("w:type"), "dxa")
    grid = table._tbl.tblGrid
    for child in list(grid):
        grid.remove(child)
    for width in widths:
        col = OxmlElement("w:gridCol")
        col.set(qn("w:w"), str(width))
        grid.append(col)
    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            tc_w.set(qn("w:w"), str(widths[idx]))
            tc_w.set(qn("w:type"), "dxa")
            cell.width = Inches(widths[idx] / 1440)
            cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
            set_cell_margins(cell)


def add_bullet(doc, text):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text))
    return p


def add_number(doc, text):
    p = doc.add_paragraph(style="List Number")
    p.paragraph_format.left_indent = Inches(0.5)
    p.paragraph_format.first_line_indent = Inches(-0.25)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.15
    set_run_font(p.add_run(text))
    return p


def add_body(doc, text, bold_prefix=None):
    p = doc.add_paragraph()
    if bold_prefix and text.startswith(bold_prefix):
        set_run_font(p.add_run(bold_prefix), bold=True)
        set_run_font(p.add_run(text[len(bold_prefix):]))
    else:
        set_run_font(p.add_run(text))
    return p


def heading(doc, text, level=1):
    p = doc.add_paragraph(style=f"Heading {level}")
    set_run_font(p.add_run(text), size={1: 20, 2: 16, 3: 14}[level], color="000000")
    return p


def build():
    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc = Document()
    sec = doc.sections[0]
    sec.page_width = Inches(8.5)
    sec.page_height = Inches(11)
    sec.top_margin = Inches(1)
    sec.bottom_margin = Inches(1)
    sec.left_margin = Inches(1)
    sec.right_margin = Inches(1)
    sec.header_distance = Inches(0.492)
    sec.footer_distance = Inches(0.492)

    styles = doc.styles
    normal = styles["Normal"]
    normal.font.name = "Arial"
    normal.font.size = Pt(11)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.15
    for idx, (size, before, after, color) in enumerate(
        ((20, 20, 6, "000000"), (16, 18, 6, "000000"), (14, 16, 4, "434343")), 1
    ):
        style = styles[f"Heading {idx}"]
        style.font.name = "Arial"
        style.font.size = Pt(size)
        style.font.bold = False
        style.font.color.rgb = RGBColor.from_string(color)
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Arial")
        style.paragraph_format.space_before = Pt(before)
        style.paragraph_format.space_after = Pt(after)
        style.paragraph_format.keep_with_next = True

    title = doc.add_paragraph()
    title.paragraph_format.space_before = Pt(0)
    title.paragraph_format.space_after = Pt(3)
    set_run_font(title.add_run("債務整理・法テラス相談用整理書面"), size=26)
    subtitle = doc.add_paragraph()
    subtitle.paragraph_format.space_after = Pt(18)
    set_run_font(subtitle.add_run("自己破産・個人再生の検討と、双極性障害に関する資料整理"), size=12, color="555555")
    meta = doc.add_paragraph()
    set_run_font(meta.add_run("作成日：2026年8月5日　　用途：弁護士・法テラスへの相談、本人用確認資料"), size=10, color="555555")

    heading(doc, "1. 現在の状況", 1)
    table = doc.add_table(rows=0, cols=2)
    table.style = "Table Grid"
    facts = [
        ("生活状況", "一人暮らし、無職"),
        ("収入", "障害年金2級。月換算約11万円"),
        ("現在の預貯金", "全口座を確認した結果、約110万円（現金・他資産を別途確認）"),
        ("過去の預貯金", "2026年5月頃は約300万円。現在までの減少約190万円について使途整理が必要"),
        ("債務名義", "楽天カード、エポス、ライフカード"),
        ("病気", "双極性障害。現在も精神科・心療内科へ通院中"),
        ("浪費", "直近約2年間で約1,500万円と指摘されている。躁状態との関係を医学資料で確認予定"),
        ("過去の財産関係", "約2年前に遺産分割、不動産持分への仮差押え、約50万円の支払い、代償金受領と生活費等への支出あり"),
    ]
    for label, value in facts:
        cells = table.add_row().cells
        set_run_font(cells[0].paragraphs[0].add_run(label), bold=True)
        set_run_font(cells[1].paragraphs[0].add_run(value))
    set_table_geometry(table, [2300, 7060])

    heading(doc, "2. 現時点での重要な整理", 1)
    add_body(doc, "自己破産を直ちに諦める段階ではない。浪費は免責不許可事由になり得るが、それだけで当然に免責不許可となるわけではなく、管財事件で裁量免責を目指す余地がある。")
    add_body(doc, "双極性障害の診断だけで免責が決まるわけではない。浪費時期と躁状態の時期が客観的に一致するか、現在の治療・家計改善・再発防止、財産と取引履歴の完全開示、管財人への協力などが重要になる。")
    add_body(doc, "個人再生は免責不許可事由の問題を避けやすいが、障害年金月約11万円から原則3年、事情により5年の返済を継続できるかを数値で検討しなければならない。")

    heading(doc, "3. 法テラス民事法律扶助の可能性", 1)
    add_body(doc, "一人暮らしの場合、資産基準は原則180万円以下、収入基準は原則月18万2,000円以下（東京都特別区・大阪市などは20万200円以下）。現在の預貯金約110万円、障害年金月約11万円という情報だけなら、資力基準を満たす可能性が高い。")
    add_body(doc, "ただし、現金、全預貯金、有価証券、暗号資産、保険解約返戻金、換価価値のある物、不動産等を合算して確認する必要がある。また、5月から現在までに約190万円減少しているため、通帳・カード明細等により使途を説明する必要がある。")
    add_body(doc, "資力基準を満たしても、自己破産では裁量免責の見込み、個人再生では再生計画の履行可能性について審査される。正式な援助申込みを行えるか、法テラス契約弁護士に確認する。")

    heading(doc, "4. 費用の目安", 1)
    heading(doc, "自己破産", 2)
    add_bullet(doc, "相談した弁護士の提示額：弁護士費用総額90万円")
    add_bullet(doc, "管財予納金：20万～50万円と説明されている")
    add_bullet(doc, "想定総額：110万～140万円程度。印紙・郵券・官報費用等が別途か、見積書で確認する")
    add_body(doc, "90万円が免責確定までの全弁護士費用であっても高額な部類であり、複雑事件としての内訳、追加費用、途中終了時の精算、申立予定日を確認する。")

    heading(doc, "個人再生", 2)
    add_bullet(doc, "弁護士費用の一般的な目安：30万～60万円程度")
    add_bullet(doc, "裁判所費用：申立手数料、郵便料、官報費用等。個人再生委員が選任される地域では約20万～30万円程度が加わることがある")
    add_bullet(doc, "概算総額：32万～90万円程度。ただし事件の複雑さ、地域、住宅ローン特則、再生委員選任の有無で変動する")
    add_body(doc, "依頼前に、弁護士費用、成功報酬、裁判所実費、個人再生委員費用を分けた見積書を受け取る。また、自己破産から個人再生へ変更する場合、または個人再生不成立後に破産へ変更する場合の追加費用を確認する。")

    heading(doc, "5. 主治医に相談する医学資料", 1)
    add_body(doc, "主治医に『免責を認めるべき』という法律判断を求めるものではない。弁護士が必要項目を整理したうえで、診療記録に基づく医学的事実を記載してもらう。いきなり意見書を依頼するのではなく、まず既存資料で足りるかを弁護士へ確認する。")
    add_bullet(doc, "障害年金申請時・更新時の診断書、精神障害者保健福祉手帳の診断書")
    add_bullet(doc, "通院歴、処方歴、カルテ・診療情報")
    add_bullet(doc, "浪費した時期に躁状態が認められたか、その時期と症状")
    add_bullet(doc, "衝動性、判断力、金銭管理能力への医学的影響")
    add_bullet(doc, "現在の病状、治療状況、再発防止に必要な支援")

    heading(doc, "主治医への依頼文例", 2)
    quote = doc.add_paragraph()
    quote.paragraph_format.left_indent = Inches(0.35)
    quote.paragraph_format.right_indent = Inches(0.35)
    quote.paragraph_format.space_before = Pt(6)
    quote.paragraph_format.space_after = Pt(10)
    set_run_font(quote.add_run("自己破産手続で、浪費した時期の病状を医学的に説明する必要があります。法的な結論ではなく、診療記録に基づき、当時の躁状態、衝動性や判断力への影響、治療経過、今後必要な再発防止策について、診断書または意見書を作成していただくことは可能でしょうか。必要であれば、弁護士から質問事項をお渡しします。"))

    heading(doc, "6. 次の弁護士・法テラスへ伝える内容", 1)
    add_body(doc, "一人暮らし、無職、障害年金2級で月約11万円、現在の現金・預貯金は約110万円です。双極性障害の躁状態と浪費時期が重なっている可能性があります。直近2年間の浪費額が大きいため、管財事件として裁量免責を目指せるか、診療記録に基づいて検討してほしいです。楽天カード、エポス、ライフカードが債務名義を取得済みです。法テラスの代理援助を正式に申し込めるか確認してください。")

    heading(doc, "確認質問", 2)
    questions = [
        "現在の全資産を約110万円とした場合、法テラスの資力基準を満たすか。正式な代理援助申込みを提出できるか。",
        "自己破産で裁量免責を目指す具体的な見通しと、免責不許可となるリスクはどの程度か。",
        "主治医の既存診断書・カルテで足りるか。追加意見書が必要なら、弁護士から質問事項を作成できるか。",
        "個人再生の場合、現在の確定債務総額、清算価値、最低弁済総額、月額、返済期間を具体的に計算してほしい。",
        "障害年金月約11万円と家賃・医療費・生活費から、個人再生を3～5年履行できるか。",
        "小規模個人再生で債権者不同意となる可能性、給与所得者等再生を利用できる可能性はあるか。",
        "契約後いつ受任通知を発送し、いつまでに申立てる予定か。債務名義3件による差押えにどう対応するか。",
        "自己破産・個人再生それぞれの費用総額、追加費用、変更時の精算を見積書で示してほしい。",
    ]
    for q in questions:
        add_number(doc, q)

    heading(doc, "7. 5月以降の約190万円の使途整理", 1)
    add_body(doc, "全口座の5月以降の明細を取得し、次の区分で集計する。分からない支出は推測で埋めず、不明として弁護士へ相談する。")
    for item in (
        "家賃・光熱費・通信費",
        "医療費・薬代・通院交通費",
        "食費・日用品など通常生活費",
        "カード会社その他債権者への返済",
        "買物、通販、ゲーム、趣味、外食など",
        "現金引出しと、その後の具体的な使途",
        "家族・知人への送金や贈与",
        "暗号資産、電子マネー、証券口座等への移動",
        "現在も手元にある購入物と、おおよその売却価値",
    ):
        add_bullet(doc, item)

    heading(doc, "8. 今後の注意事項", 1)
    for item in (
        "新たな借入れやカード利用をしない。",
        "家族・知人への送金、財産の名義変更、多額の現金引出しをしない。",
        "楽天・エポス・ライフカードなど、特定の債権者だけに返済しない。",
        "通帳、カード明細、通販履歴、領収書、診療記録を保存する。",
        "家計表を毎月作成し、通常生活費も記録する。",
        "通院・服薬を継続し、カード利用停止や金銭管理支援など再発防止策を相談する。",
        "弁護士、法テラス、裁判所、破産管財人に事実を隠さず、分からないことは分からないと説明する。",
        "私費契約・90万円の支払い前に、法テラス利用とセカンドオピニオンを検討する。",
    ):
        add_bullet(doc, item)

    heading(doc, "9. 参考情報", 1)
    refs = [
        "法テラス『弁護士等費用の立替制度のご利用の流れ』 https://www.houterasu.or.jp/site/soudan-tatekae/goriyou.html",
        "法テラス『民事法律扶助業務』 https://www.houterasu.or.jp/site/bengoshitou-fujo/index.html",
        "裁判所『個人再生』 https://www.courts.go.jp/saiban/syurui/syurui_minzi/minzi_25_18/index.html",
        "大阪地方裁判所『倒産部（破産・個人再生Q&A）』 https://www.courts.go.jp/osaka/saiban/minjibu6/index.html",
        "千葉県弁護士会『個人の債務整理に関する弁護士費用の目安』 https://www.chiba-ben.or.jp/soudan/consultation/syakkin_hiyou.html",
    ]
    for ref in refs:
        add_bullet(doc, ref)
    note = doc.add_paragraph()
    note.paragraph_format.space_before = Pt(12)
    set_run_font(note.add_run("注意：本書面は相談内容の整理用であり、個別事件の法的結論を確定するものではありません。最終判断は、資料を確認した弁護士と管轄裁判所によります。"), size=9, color="555555")

    doc.core_properties.title = "債務整理・法テラス相談用整理書面"
    doc.core_properties.subject = "自己破産、個人再生、双極性障害に関する相談整理"
    doc.core_properties.author = ""
    doc.core_properties.keywords = ""
    doc.save(OUT)
    print(OUT)


if __name__ == "__main__":
    build()
